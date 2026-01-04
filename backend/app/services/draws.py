from __future__ import annotations

import random
from io import BytesIO

from fastapi import HTTPException, status
from openpyxl import Workbook
from sqlalchemy import and_, delete, func, or_, select
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy.orm import Session, selectinload

from app.models.draw import DrawApplication, DrawResult
from app.models.expert import Expert
from app.models.expert_specialty import ExpertSpecialty
from app.models.rule import Rule
from app.models.organization import Organization
from app.models.specialty import Specialty
from app.repo.draws import DrawRepo
from app.repo.rules import RuleRepo
from app.repo.utils import apply_keyword, apply_sort, paginate
from app.services import experts as expert_service
from app.services import specialties as specialty_service
from app.services import titles as title_service
from app.schemas.pagination import PageParams
from app.schemas.draw import DrawApply, DrawUpdate

SUPPORTED_DRAW_METHODS = {"random", "lottery"}
CONTACT_STATUS_PENDING = "pending"
CONTACT_STATUS_ACCEPTED = "accepted"
CONTACT_STATUS_REJECTED = "rejected"


def _split_terms(value: str | None) -> list[str]:
    if not value:
        return []
    normalized = value
    for sep in ["、", "，", "；", ",", ";", "|", "\n"]:
        normalized = normalized.replace(sep, ";")
    return [item.strip() for item in normalized.split(";") if item.strip()]


def _split_numeric_terms(value: str | None) -> tuple[list[int], list[str]]:
    numeric: list[int] = []
    text: list[str] = []
    for item in _split_terms(value):
        if item.isdigit():
            numeric.append(int(item))
        else:
            text.append(item)
    return numeric, text


def _unique_ints(values: list[int] | None) -> list[int]:
    unique: list[int] = []
    for item in values or []:
        try:
            value = int(item)
        except (TypeError, ValueError):
            continue
        if value not in unique:
            unique.append(value)
    return unique


def _split_person_terms(
    value: str | None,
) -> tuple[list[int], list[str]]:
    expert_ids: list[int] = []
    invalid: list[str] = []
    for item in _split_terms(value):
        normalized = item.strip()
        if not normalized:
            continue
        if normalized.isdigit():
            expert_ids.append(int(normalized))
            continue
        invalid.append(normalized)
    return expert_ids, invalid


def _mask_phone(value: str | None) -> str | None:
    if not value:
        return None
    raw = str(value).strip()
    if not raw:
        return None
    if len(raw) <= 7:
        return raw
    return f"{raw[:3]}{'*' * (len(raw) - 7)}{raw[-4:]}"


def _mask_id_card(value: str | None) -> str | None:
    if not value:
        return None
    raw = str(value).strip()
    if not raw:
        return None
    if len(raw) <= 7:
        return raw
    return f"{raw[:3]}{'*' * (len(raw) - 7)}{raw[-4:]}"


def _mask_name(value: str | None) -> str | None:
    if not value:
        return None
    raw = str(value).strip()
    if not raw:
        return None
    if len(raw) == 1:
        return raw
    if len(raw) == 2:
        return f"{raw[0]}*"
    return f"{raw[0]}{'*' * (len(raw) - 2)}{raw[-1]}"


def _find_active_rule_by_specialties(
    db: Session, specialty_ids: list[int], specialty_name: str | None
) -> Rule | None:
    if not specialty_ids and not specialty_name:
        return None
    stmt = select(Rule).where(Rule.is_active.is_(True)).order_by(Rule.id.desc())
    rules = db.execute(stmt).scalars().all()
    name_terms = _split_terms(specialty_name) if specialty_name else []
    specialty_id_set = set(specialty_ids)
    for rule in rules:
        rule_ids = _unique_ints(rule.specialty_ids)
        if specialty_id_set and rule_ids:
            if specialty_id_set.intersection(rule_ids):
                return rule
        if name_terms:
            rule_names = _split_terms(rule.specialty)
            if any(name in rule_names for name in name_terms):
                return rule
    return None


def resolve_draw_method(draw: DrawApplication, rule) -> str:
    if draw.draw_method:
        return draw.draw_method
    if rule is not None and rule.draw_method:
        return rule.draw_method
    return "random"


def _resolve_total_count(expert_count: int, total_count: int | None) -> int:
    resolved = total_count if total_count is not None else expert_count
    if resolved < expert_count:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="抽取人数不能大于总人数",
        )
    return resolved


def pick_experts(
    candidates: list[Expert], total_needed: int, draw_method: str
) -> list[Expert]:
    if draw_method == "random":
        return random.sample(candidates, total_needed)
    if draw_method == "lottery":
        tickets = [(random.random(), expert) for expert in candidates]
        tickets.sort(key=lambda item: item[0])
        return [expert for _, expert in tickets[:total_needed]]
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Unsupported draw method",
    )


def list_draws(db: Session, params: PageParams) -> tuple[list[DrawApplication], int]:
    return DrawRepo(db).list_page(
        params.keyword,
        params.sort_by,
        params.sort_order,
        params.page,
        params.page_size,
    )


def get_draw(db: Session, draw_id: int) -> DrawApplication:
    draw = DrawRepo(db).get_by_id(draw_id)
    if draw is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Draw not found")
    return draw


def create_draw(db: Session, payload: DrawApply, created_by_id: int | None) -> DrawApplication:
    data = payload.model_dump()
    rule_id = data.get("rule_id")
    if rule_id is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rule is required",
        )
    rule = RuleRepo(db).get_by_id(rule_id)
    if rule is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Rule not found"
        )

    expert_count = data.get("expert_count") or 0
    data["total_count"] = _resolve_total_count(
        expert_count, data.get("total_count")
    )
    draw = DrawApplication(**data, created_by_id=created_by_id)
    draw.rule_id = rule.id
    draw.category_id = None
    draw.category = rule.category or "不限"
    draw.subcategory_id = None
    draw.subcategory = None
    draw.specialty_id = rule.specialty_id
    draw.specialty = rule.specialty

    db.add(draw)
    db.commit()
    db.refresh(draw)
    return draw


def update_draw(db: Session, draw_id: int, payload: DrawUpdate) -> DrawApplication:
    draw = get_draw(db, draw_id)
    update_data = payload.model_dump(exclude_unset=True)
    condition_fields = {
        "rule_id",
        "expert_count",
        "backup_count",
        "draw_method",
        "avoid_units",
        "avoid_persons",
    }
    original_conditions = {field: getattr(draw, field) for field in condition_fields}
    original_status = draw.status

    if "rule_id" in update_data:
        rule_id = update_data.get("rule_id")
        if rule_id is None:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Rule is required",
            )
        rule = RuleRepo(db).get_by_id(rule_id)
        if rule is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND, detail="Rule not found"
            )
        draw.rule_id = rule.id
        draw.category_id = None
        draw.category = rule.category or "不限"
        draw.subcategory_id = None
        draw.subcategory = None
        draw.specialty_id = rule.specialty_id
        draw.specialty = rule.specialty

    for key, value in update_data.items():
        if key in {"rule_id"}:
            continue
        if value is None and key in {"expert_count", "total_count"}:
            continue
        setattr(draw, key, value)

    new_expert_count = draw.expert_count
    new_total_count = draw.total_count
    if "expert_count" in update_data and update_data.get("expert_count") is not None:
        new_expert_count = update_data.get("expert_count") or 0
    if "total_count" in update_data:
        new_total_count = update_data.get("total_count")
    draw.total_count = _resolve_total_count(new_expert_count, new_total_count)

    conditions_changed = any(
        getattr(draw, field) != original_conditions[field]
        for field in condition_fields
    )
    status_changed = "status" in update_data and update_data["status"] != original_status
    reset_results = conditions_changed
    if status_changed and update_data["status"] in {"pending", "scheduled"}:
        reset_results = True
    if reset_results:
        db.execute(delete(DrawResult).where(DrawResult.draw_id == draw_id))
        if "status" not in update_data and draw.status != "cancelled":
            draw.status = "pending"

    db.commit()
    db.refresh(draw)
    return draw


def delete_draw(db: Session, draw_id: int) -> None:
    draw = get_draw(db, draw_id)
    db.delete(draw)
    db.commit()


def delete_draws(db: Session, draw_ids: list[int]) -> dict[str, int]:
    unique_ids = {int(item) for item in draw_ids if isinstance(item, int)}
    if not unique_ids:
        return {"deleted": 0, "skipped": 0}
    existing = set(
        db.execute(select(DrawApplication.id).where(DrawApplication.id.in_(unique_ids)))
        .scalars()
        .all()
    )
    if not existing:
        return {"deleted": 0, "skipped": len(unique_ids)}
    db.execute(delete(DrawResult).where(DrawResult.draw_id.in_(existing)))
    db.execute(delete(DrawApplication).where(DrawApplication.id.in_(existing)))
    db.commit()
    return {"deleted": len(existing), "skipped": len(unique_ids) - len(existing)}


def list_results(db: Session, draw_id: int) -> list[DrawResult]:
    _ = get_draw(db, draw_id)
    stmt = (
        select(DrawResult)
        .where(DrawResult.draw_id == draw_id)
        .options(selectinload(DrawResult.expert))
        .order_by(DrawResult.is_backup, DrawResult.ordinal)
    )
    results = list(db.execute(stmt).scalars().all())
    experts = [result.expert for result in results if result.expert]
    expert_service._attach_expert_details(db, experts)
    return results


def export_results(db: Session, draw_id: int) -> BytesIO:
    draw = get_draw(db, draw_id)
    results = list_results(db, draw_id)
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.append(
        [
            "抽取编号",
            "项目名称",
            "项目编号",
            "序号",
            "候补",
            "递补",
            "专家姓名",
            "单位",
            "专业",
            "职称",
            "电话",
            "身份证号",
        ]
    )

    for result in results:
        expert = result.expert
        specialties = getattr(expert, "specialties", []) if expert else []
        specialty_names = ";".join([item.name for item in specialties if item.name])
        worksheet.append(
            [
                draw.id,
                draw.project_name or "",
                draw.project_code or "",
                result.ordinal or "",
                "是" if result.is_backup else "否",
                "是" if result.is_replacement else "否",
                _mask_name(expert.name) if expert else "",
                expert.company if expert else "",
                specialty_names,
                expert.title if expert else "",
                _mask_phone(expert.phone) if expert else "",
                _mask_id_card(expert.id_card_no) if expert else "",
            ]
        )

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def _format_review_time(value) -> str:
    if not value:
        return ""
    try:
        return value.strftime("%Y-%m-%d %H:%M")
    except Exception:
        return str(value)


def export_signin_sheet(db: Session, draw_id: int) -> BytesIO:
    draw = get_draw(db, draw_id)
    results = list_results(db, draw_id)
    main_results = [item for item in results if not item.is_backup]
    main_results.sort(key=lambda item: (item.ordinal or 0, item.id))

    total_count = draw.total_count or draw.expert_count or len(main_results)
    if total_count < len(main_results):
        total_count = len(main_results)

    avoid_unit_ids, avoid_unit_names = _split_numeric_terms(draw.avoid_units)
    if avoid_unit_ids:
        names = (
            db.execute(
                select(Organization.name).where(Organization.id.in_(avoid_unit_ids))
            )
            .scalars()
            .all()
        )
        avoid_unit_names.extend(names)
    seen_units: list[str] = []
    for name in avoid_unit_names:
        if name not in seen_units:
            seen_units.append(name)
    avoid_units_label = "；".join(seen_units)

    specialty_label = draw.specialty or draw.category or ""

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "FangSong"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    normal_style.font.size = Pt(10.5)

    title = doc.add_paragraph(
        "云南嘉顺工程项目管理有限公司评审专家抽取结果暨评审组成员签到表"
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(16)

    headers = ["序号", "评审人员", "工作单位", "身份证号码", "手机号码", "签名"]
    row_count = 9 + total_count
    table = doc.add_table(rows=row_count, cols=6)
    table.style = "Table Grid"
    col_widths = [Mm(20), Mm(28), Mm(68), Mm(58), Mm(35), Mm(20)]
    for column, width in zip(table.columns, col_widths):
        for cell in column.cells:
            cell.width = width

    def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=None):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        run = paragraph.add_run(text)
        run.bold = bold
        if size:
            run.font.size = size

    def clear_cell_border(cell) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = tc_borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tc_borders.append(element)
            element.set(qn("w:val"), "nil")

    # Row 0: purchaser
    table.cell(0, 1).merge(table.cell(0, 5))
    set_cell_text(table.cell(0, 0), "采购单位")
    set_cell_text(table.cell(0, 1), "")

    # Row 1: project name/code
    table.cell(1, 1).merge(table.cell(1, 3))
    set_cell_text(table.cell(1, 0), "项目名称")
    set_cell_text(table.cell(1, 1), draw.project_name or "")
    set_cell_text(table.cell(1, 4), "项目编号")
    set_cell_text(table.cell(1, 5), draw.project_code or "")

    # Row 2: review location/time
    table.cell(2, 1).merge(table.cell(2, 3))
    set_cell_text(table.cell(2, 0), "评审地点")
    set_cell_text(table.cell(2, 1), draw.review_location or "")
    set_cell_text(table.cell(2, 4), "评审开始时间")
    set_cell_text(table.cell(2, 5), _format_review_time(draw.review_time))

    # Row 3: specialty
    table.cell(3, 1).merge(table.cell(3, 5))
    set_cell_text(table.cell(3, 0), "专业类别")
    set_cell_text(table.cell(3, 1), specialty_label)

    # Row 4: avoid units
    table.cell(4, 1).merge(table.cell(4, 5))
    set_cell_text(table.cell(4, 0), "屏蔽单位")
    set_cell_text(table.cell(4, 1), avoid_units_label)

    # Row 5: totals (4 cells)
    table.cell(5, 0).merge(table.cell(5, 1))
    table.cell(5, 3).merge(table.cell(5, 4))
    set_cell_text(table.cell(5, 0), "总人数")
    set_cell_text(table.cell(5, 2), str(total_count))
    set_cell_text(table.cell(5, 3), "抽取人数")
    set_cell_text(table.cell(5, 5), str(draw.expert_count))

    # Row 6: group label
    table.cell(6, 0).merge(table.cell(6, 5))
    set_cell_text(table.cell(6, 0), "评审组成员：")

    # Row 7: header
    for idx, text in enumerate(headers):
        set_cell_text(table.cell(7, idx), text)

    # Rows 8..: experts
    start_row = 8
    list_height = Mm(8)
    for row in table.rows[7 : start_row + total_count]:
        row.height = list_height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    meta_height = Mm(8)
    for row in table.rows[:7]:
        row.height = meta_height
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for index in range(total_count):
        row = table.rows[start_row + index].cells
        set_cell_text(row[0], str(index + 1))
        if index < len(main_results):
            expert = main_results[index].expert
            set_cell_text(row[1], expert.name if expert else "")
            set_cell_text(row[2], expert.company if expert else "")
            set_cell_text(row[3], expert.id_card_no if expert else "")
            set_cell_text(row[4], expert.phone if expert else "")
        set_cell_text(row[5], "")

    # Final row: draw/supervisor
    footer_row = table.rows[start_row + total_count].cells
    footer_row[0].merge(footer_row[2])
    footer_row[3].merge(footer_row[5])
    set_cell_text(footer_row[0], "抽取人员：", bold=True)
    set_cell_text(footer_row[3], "监督人员：", bold=True)
    footer = table.rows[start_row + total_count]
    footer.height = list_height
    footer.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for cell in footer.cells:
        clear_cell_border(cell)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def list_results_page(
    db: Session, draw_id: int, params: PageParams
) -> tuple[list[DrawResult], int]:
    _ = get_draw(db, draw_id)
    stmt = (
        select(DrawResult)
        .where(DrawResult.draw_id == draw_id)
        .options(selectinload(DrawResult.expert))
    )
    if params.keyword:
        stmt = stmt.join(Expert, Expert.id == DrawResult.expert_id, isouter=True)
        stmt = apply_keyword(
            stmt,
            params.keyword,
            [Expert.name, Expert.company, Expert.phone],
        )
    sort_map = {
        "id": DrawResult.id,
        "ordinal": DrawResult.ordinal,
        "is_backup": DrawResult.is_backup,
        "is_replacement": DrawResult.is_replacement,
    }
    if params.sort_by:
        stmt = apply_sort(
            stmt, params.sort_by, params.sort_order, sort_map, DrawResult.id
        )
    else:
        stmt = stmt.order_by(DrawResult.is_backup, DrawResult.ordinal, DrawResult.id)
    items, total = paginate(db, stmt, params.page, params.page_size)
    experts = [result.expert for result in items if result.expert]
    expert_service._attach_expert_details(db, experts)
    return items, total


def execute_draw(db: Session, draw_id: int) -> list[DrawResult]:
    draw = get_draw(db, draw_id)
    if draw.status == "cancelled":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Draw already completed or cancelled",
        )

    existing_results = (
        db.execute(select(DrawResult).where(DrawResult.draw_id == draw.id))
        .scalars()
        .all()
    )
    if existing_results:
        has_confirmed = any(
            (
                result.contact_status
                in {
                    CONTACT_STATUS_ACCEPTED,
                    CONTACT_STATUS_REJECTED,
                }
            )
            for result in existing_results
            if not result.is_backup
        )
        updated = False
        if draw.status == "completed":
            for result in existing_results:
                if result.is_backup:
                    continue
                if result.contact_status is None:
                    result.contact_status = CONTACT_STATUS_ACCEPTED
                    updated = True
        else:
            if has_confirmed:
                _apply_completion_status(db, draw)
                updated = True
            else:
                if draw.status != "scheduled":
                    draw.status = "scheduled"
                    updated = True
                for result in existing_results:
                    if result.contact_status is None:
                        result.contact_status = CONTACT_STATUS_PENDING
                        updated = True
        if updated:
            db.commit()
        return list_results(db, draw.id)
    if draw.status == "completed":
        return list_results(db, draw.id)

    rule = RuleRepo(db).get_by_id(draw.rule_id) if draw.rule_id else None
    if rule is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rule is required",
        )

    stmt = select(Expert).where(Expert.is_active.is_(True)).distinct()
    specialty_ids = specialty_service.expand_to_leaf_ids(
        db, _unique_ints(rule.specialty_ids)
    )
    specialty_names = _split_terms(rule.specialty) if not specialty_ids else []
    if specialty_ids or specialty_names:
        stmt = stmt.join(
            ExpertSpecialty, ExpertSpecialty.expert_id == Expert.id
        ).join(Specialty, Specialty.id == ExpertSpecialty.specialty_id)
        if specialty_ids:
            stmt = stmt.where(Specialty.id.in_(specialty_ids))
        else:
            stmt = stmt.where(Specialty.name.in_(specialty_names))

    title_required_ids = title_service.expand_to_leaf_ids(
        db, _unique_ints(rule.title_required_ids)
    )
    title_names = _split_terms(rule.title_required)
    if title_required_ids:
        if title_names:
            stmt = stmt.where(
                or_(
                    Expert.title_id.in_(title_required_ids),
                    Expert.title.in_(title_names),
                    (Expert.title_id.is_(None) & Expert.title.is_(None)),
                )
            )
        else:
            stmt = stmt.where(
                or_(
                    Expert.title_id.in_(title_required_ids),
                    (Expert.title_id.is_(None) & Expert.title.is_(None)),
                )
            )
    elif rule.title_required:
        stmt = stmt.where(
            or_(
                Expert.title.in_(title_names),
                (Expert.title_id.is_(None) & Expert.title.is_(None)),
            )
        )

    region_required_ids = _unique_ints(rule.region_required_ids)
    region_names = _split_terms(rule.region_required)
    if region_required_ids:
        if region_names:
            stmt = stmt.where(
                or_(
                    Expert.region_id.in_(region_required_ids),
                    Expert.region.in_(region_names),
                )
            )
        else:
            stmt = stmt.where(Expert.region_id.in_(region_required_ids))
    elif rule.region_required_id is not None:
        stmt = stmt.where(Expert.region_id == rule.region_required_id)
    elif rule.region_required:
        stmt = stmt.where(Expert.region.in_(region_names))
    avoid_unit_ids, avoid_unit_names = _split_numeric_terms(draw.avoid_units)
    if avoid_unit_names:
        matched_ids = (
            db.execute(
                select(Organization.id).where(Organization.name.in_(avoid_unit_names))
            )
            .scalars()
            .all()
        )
        for org_id in matched_ids:
            if org_id not in avoid_unit_ids:
                avoid_unit_ids.append(org_id)
    unit_filters = []
    if avoid_unit_ids:
        unit_filters.append(
            or_(
                Expert.organization_id.is_(None),
                ~Expert.organization_id.in_(avoid_unit_ids),
            )
        )
    if avoid_unit_names:
        unit_filters.append(
            or_(
                Expert.company.is_(None),
                ~Expert.company.in_(avoid_unit_names),
            )
        )
    if unit_filters:
        stmt = stmt.where(and_(*unit_filters))

    avoid_person_ids, invalid_person_terms = _split_person_terms(draw.avoid_persons)
    if invalid_person_terms:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="回避人员必须选择专家",
        )
    if avoid_person_ids:
        existing = set(
            db.execute(
                select(Expert.id).where(Expert.id.in_(avoid_person_ids))
            )
            .scalars()
            .all()
        )
        missing = [item for item in avoid_person_ids if item not in existing]
        if missing:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="回避人员不存在",
            )
        stmt = stmt.where(~Expert.id.in_(avoid_person_ids))

    candidates = list(db.execute(stmt).scalars().all())
    backup_count = draw.backup_count or 0
    total_needed = draw.expert_count + backup_count
    if len(candidates) < total_needed:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Not enough qualified experts",
        )

    method = resolve_draw_method(draw, rule)
    if method not in SUPPORTED_DRAW_METHODS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported draw method",
        )
    chosen = pick_experts(candidates, total_needed, method)
    for index, expert in enumerate(chosen, start=1):
        is_backup = index > draw.expert_count
        result = DrawResult(
            draw_id=draw.id,
            expert_id=expert.id,
            is_backup=is_backup,
            contact_status=CONTACT_STATUS_PENDING,
            ordinal=index,
        )
        db.add(result)

    draw.draw_method = method
    draw.status = "scheduled"
    db.commit()

    return list_results(db, draw.id)


def _apply_completion_status(db: Session, draw: DrawApplication) -> None:
    if draw.status == "cancelled":
        return
    db.flush()
    accepted_count = (
        db.execute(
            select(func.count())
            .select_from(DrawResult)
            .where(
                DrawResult.draw_id == draw.id,
                DrawResult.is_backup.is_(False),
                DrawResult.contact_status == CONTACT_STATUS_ACCEPTED,
            )
        )
        .scalar_one()
    )
    if accepted_count >= draw.expert_count:
        draw.status = "completed"
    else:
        draw.status = "scheduled"


def replace_draw_result(db: Session, draw_id: int, result_id: int) -> list[DrawResult]:
    draw = get_draw(db, draw_id)
    if draw.status == "cancelled":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Draw already completed or cancelled",
        )

    target = db.execute(
        select(DrawResult).where(
            DrawResult.draw_id == draw.id, DrawResult.id == result_id
        )
    ).scalar_one_or_none()
    if target is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Result not found"
        )
    if target.is_backup:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot replace a backup expert",
        )

    backup = (
        db.execute(
            select(DrawResult)
            .where(DrawResult.draw_id == draw.id, DrawResult.is_backup.is_(True))
            .order_by(DrawResult.ordinal, DrawResult.id)
        )
        .scalars()
        .first()
    )
    if backup is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No backup experts available",
        )

    target_ordinal = target.ordinal
    db.delete(target)
    backup.is_backup = False
    backup.is_replacement = True
    backup.contact_status = CONTACT_STATUS_PENDING
    if target_ordinal is not None:
        backup.ordinal = target_ordinal
    _apply_completion_status(db, draw)
    db.commit()

    return list_results(db, draw.id)


def get_draw_result_contact(
    db: Session, draw_id: int, result_id: int
) -> dict[str, str | None]:
    _ = get_draw(db, draw_id)
    result = (
        db.execute(
            select(DrawResult)
            .where(DrawResult.draw_id == draw_id, DrawResult.id == result_id)
            .options(selectinload(DrawResult.expert))
        )
        .scalars()
        .first()
    )
    if result is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Result not found"
        )
    if result.is_backup:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="候补专家无需联系",
        )
    expert = result.expert
    if expert is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Expert not found"
        )
    return {"name": expert.name, "phone": expert.phone}


def update_draw_result_contact(
    db: Session,
    draw_id: int,
    result_id: int,
    status_value: str,
    auto_replace: bool,
) -> list[DrawResult]:
    draw = get_draw(db, draw_id)
    if draw.status == "cancelled":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Draw already completed or cancelled",
        )
    result = db.execute(
        select(DrawResult).where(
            DrawResult.draw_id == draw.id, DrawResult.id == result_id
        )
    ).scalar_one_or_none()
    if result is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Result not found"
        )
    if result.is_backup:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="候补专家无需联系",
        )
    if status_value not in {
        CONTACT_STATUS_ACCEPTED,
        CONTACT_STATUS_REJECTED,
        CONTACT_STATUS_PENDING,
    }:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="联系结果无效",
        )

    if status_value == CONTACT_STATUS_REJECTED and auto_replace:
        return replace_draw_result(db, draw_id, result_id)

    result.contact_status = status_value
    _apply_completion_status(db, draw)
    db.commit()
    return list_results(db, draw.id)
